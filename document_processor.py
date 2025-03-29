import os
from typing import Tuple, Dict, Optional
import PyPDF2
import docx
import re

def process_document(file_path: str, original_filename: str) -> Tuple[str, Dict]:
    """
    Process different document formats and extract text content
    
    Args:
        file_path: Path to the temporary file
        original_filename: Original name of the uploaded file
        
    Returns:
        Tuple containing extracted text content and metadata dictionary
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        return extract_from_pdf(file_path)
    elif file_extension == '.docx':
        return extract_from_docx(file_path)
    elif file_extension == '.txt':
        return extract_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")

def extract_from_pdf(file_path: str) -> Tuple[str, Dict]:
    """Extract text and metadata from PDF file"""
    content = ""
    metadata = {}
    
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Extract metadata if available
            if pdf_reader.metadata:
                for key, value in pdf_reader.metadata.items():
                    if key.startswith('/'):
                        clean_key = key[1:]  # Remove leading slash
                        metadata[clean_key] = value
            
            # Extract text from each page
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                content += page.extract_text() + "\n\n"
            
            # Add page count to metadata
            metadata['PageCount'] = len(pdf_reader.pages)
            
    except Exception as e:
        raise Exception(f"Error processing PDF: {str(e)}")
    
    return clean_text(content), metadata

def extract_from_docx(file_path: str) -> Tuple[str, Dict]:
    """Extract text and metadata from DOCX file"""
    content = ""
    metadata = {
        'PageCount': 0,
        'Paragraphs': 0
    }
    
    try:
        doc = docx.Document(file_path)
        
        # Extract metadata from core properties if available
        if hasattr(doc, 'core_properties'):
            props = doc.core_properties
            if props.title:
                metadata['Title'] = props.title
            if props.author:
                metadata['Author'] = props.author
            if props.created:
                metadata['Created'] = str(props.created)
            if props.modified:
                metadata['Modified'] = str(props.modified)
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                content += para.text + "\n\n"
                metadata['Paragraphs'] += 1
        
        # Handle tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([cell.text for cell in row.cells])
                content += row_text + "\n"
        
    except Exception as e:
        raise Exception(f"Error processing DOCX: {str(e)}")
    
    return clean_text(content), metadata

def extract_from_txt(file_path: str) -> Tuple[str, Dict]:
    """Extract text from TXT file"""
    content = ""
    metadata = {
        'FileSize': os.path.getsize(file_path),
        'LineCount': 0
    }
    
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
            lines = file.readlines()
            metadata['LineCount'] = len(lines)
            content = ''.join(lines)
    except Exception as e:
        # Try with different encodings if UTF-8 fails
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                lines = file.readlines()
                metadata['LineCount'] = len(lines)
                content = ''.join(lines)
                metadata['Encoding'] = 'latin-1'
        except Exception as e2:
            raise Exception(f"Error processing TXT: {str(e)} and {str(e2)}")
    
    return clean_text(content), metadata

def clean_text(text: str) -> str:
    """Clean and normalize text content"""
    if not text:
        return ""
    
    # Replace multiple whitespaces with a single space
    text = re.sub(r'\s+', ' ', text)
    
    # Remove non-breaking spaces and other special whitespace
    text = text.replace('\xa0', ' ').strip()
    
    # Remove any control characters
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    
    return text
